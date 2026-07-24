"""
Document parser using Docling + RapidOCR for structured Markdown output.

Responsibilities:
  - Validate file type and size
  - Route to the correct parser based on file extension
  - Return extracted text with metadata

Supported formats:
  - .pdf/.png/.jpg/.jpeg/.tiff → Docling (text + scanned via RapidOCR)
  - .docx                     → python-docx
  - .txt/.md                  → plain UTF-8
  - .csv                      → csv module
  - .json                     → json module
  - .html                     → html.parser
"""
import csv
import io
import json
import logging
import os
import tempfile
from html.parser import HTMLParser
from pathlib import Path

logger = logging.getLogger(__name__)


# ─── Constants ────────────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {
    ".pdf", ".png", ".jpg", ".jpeg", ".tiff",
    ".docx",
    ".txt", ".md", ".csv", ".json", ".html",
}

DOCLING_EXTENSIONS = {".pdf", ".png", ".jpg", ".jpeg", ".tiff"}

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB


# ─── Validation ───────────────────────────────────────────────────────────────

def _validate_file_size(file_size: int) -> None:
    """Raise ValueError if file exceeds the max allowed size."""
    if file_size > MAX_FILE_SIZE_BYTES:
        size_mb = file_size / (1024 * 1024)
        max_mb = MAX_FILE_SIZE_BYTES / (1024 * 1024)
        logger.warning("File rejected: %.1fMB exceeds %.0fMB limit", size_mb, max_mb)
        raise ValueError(f"File too large: {size_mb:.1f}MB. Maximum is {max_mb:.0f}MB.")


def _validate_extension(ext: str) -> None:
    """Raise ValueError if the file extension is not supported."""
    if ext not in SUPPORTED_EXTENSIONS:
        logger.warning("Unsupported file extension: '%s'", ext)
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )


# ─── Docling Converter ───────────────────────────────────────────────────────

class _DoclingManager:
    """Lazy-loads and caches the Docling DocumentConverter singleton."""

    def __init__(self):
        self._converter = None

    def get_converter(self):
        """Return the Docling converter, initializing it on first call."""
        if self._converter is not None:
            return self._converter

        try:
            from docling.document_converter import DocumentConverter, PdfFormatOption
            from docling.datamodel.pipeline_options import PdfPipelineOptions, RapidOcrOptions
            from docling.datamodel.base_models import InputFormat

            pipeline_options = PdfPipelineOptions()
            pipeline_options.do_ocr = True
            pipeline_options.ocr_options = RapidOcrOptions()

            self._converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options),
                }
            )
            logger.info("Docling converter initialized with RapidOCR backend")
        except ImportError as e:
            logger.warning("Docling not available: %s. PDF/image parsing will use fallbacks.", e)
            self._converter = None

        return self._converter


_docling_manager = _DoclingManager()


# ─── Individual Parsers ───────────────────────────────────────────────────────

def _parse_with_docling(content: bytes, ext: str) -> tuple[str, str]:
    """Use Docling to convert PDF/image → structured Markdown with RapidOCR."""
    converter = _docling_manager.get_converter()

    if converter is None:
        if ext == ".pdf":
            logger.info("Docling unavailable — falling back to pypdf for PDF")
            return _parse_pdf_fallback(content)
        logger.error("Docling unavailable — cannot parse '%s' files without it", ext)
        return "", f"error: docling not available for {ext}"

    suffix = ext if ext != ".jpeg" else ".jpg"
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
            logger.debug("Wrote %d bytes to temp file: %s", len(content), tmp_path)

        result = converter.convert(tmp_path)
        markdown_text = result.document.export_to_markdown()
        logger.info("Docling conversion succeeded — extracted %d chars from %s", len(markdown_text), ext)
        return markdown_text, "docling+rapidocr"
    except Exception as e:
        logger.error("Docling conversion failed for %s: %s", ext, e, exc_info=True)
        if ext == ".pdf":
            logger.info("Falling back to pypdf after Docling failure")
            return _parse_pdf_fallback(content)
        return f"[Error parsing file: {e}]", "error"
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                logger.debug("Could not delete temp file: %s", tmp_path)


def _parse_pdf_fallback(content: bytes) -> tuple[str, str]:
    """Fallback PDF text extraction using pypdf if docling fails."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(text)
                logger.debug("pypdf: extracted %d chars from page %d", len(text), i + 1)
        combined = "\n\n".join(pages)
        logger.info("pypdf fallback extracted %d chars from %d pages", len(combined), len(pages))
        return combined, "pypdf-fallback"
    except ImportError:
        logger.error("pypdf not installed — PDF parsing completely unavailable")
        return "[PDF parsing unavailable. Install docling or pypdf.]", "error"
    except Exception as e:
        logger.error("pypdf fallback failed: %s", e, exc_info=True)
        return f"[Error reading PDF: {e}]", "error"


def _parse_docx(content: bytes) -> tuple[str, str]:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    try:
                        level_num = int(level)
                        parts.append("#" * level_num + " " + para.text.strip())
                    except ValueError:
                        parts.append(para.text.strip())
                else:
                    parts.append(para.text.strip())

        for table in doc.tables:
            rows = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    rows.append("|" + " --- |" * len(cells))
            parts.append("\n".join(rows))

        combined = "\n\n".join(parts)
        logger.info("python-docx extracted %d chars, %d paragraphs, %d tables", len(combined), len(doc.paragraphs), len(doc.tables))
        return combined, "python-docx"
    except ImportError:
        logger.error("python-docx not installed")
        return "[DOCX parsing unavailable. Install python-docx.]", "error"
    except Exception as e:
        logger.error("DOCX parsing failed: %s", e, exc_info=True)
        return f"[Error reading DOCX: {e}]", "error"


def _parse_csv(content: bytes) -> tuple[str, str]:
    """Convert CSV to markdown table format."""
    try:
        text = content.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)

        if not rows:
            logger.info("CSV file is empty")
            return "", "csv-builtin"

        lines = []
        for i, row in enumerate(rows):
            lines.append("| " + " | ".join(cell.strip() for cell in row) + " |")
            if i == 0:
                lines.append("|" + " --- |" * len(row))

        result = "\n".join(lines)
        logger.info("CSV parsed: %d rows, %d chars output", len(rows), len(result))
        return result, "csv-builtin"
    except Exception as e:
        logger.error("CSV parsing failed: %s", e, exc_info=True)
        return f"[Error reading CSV: {e}]", "error"


def _parse_json(content: bytes) -> tuple[str, str]:
    """Pretty-print JSON content."""
    try:
        text = content.decode("utf-8", errors="replace")
        data = json.loads(text)
        result = json.dumps(data, indent=2, ensure_ascii=False)
        logger.info("JSON parsed: %d chars output", len(result))
        return result, "json-builtin"
    except json.JSONDecodeError:
        logger.warning("Invalid JSON — returning raw text content as fallback")
        return content.decode("utf-8", errors="replace"), "text-fallback"
    except Exception as e:
        logger.error("JSON parsing failed: %s", e, exc_info=True)
        return f"[Error reading JSON: {e}]", "error"


def _parse_html(content: bytes) -> tuple[str, str]:
    """Strip HTML tags and extract text."""
    class _TextExtractor(HTMLParser):
        def __init__(self):
            super().__init__()
            self.chunks: list[str] = []
            self._skip_tags = {"script", "style", "meta", "link", "head"}
            self._skip = False

        def handle_starttag(self, tag, attrs):
            if tag in self._skip_tags:
                self._skip = True

        def handle_endtag(self, tag):
            if tag in self._skip_tags:
                self._skip = False

        def handle_data(self, data):
            if not self._skip:
                data = data.strip()
                if data:
                    self.chunks.append(data)

    try:
        text = content.decode("utf-8", errors="replace")
        extractor = _TextExtractor()
        extractor.feed(text)
        result = "\n\n".join(extractor.chunks)
        logger.info("HTML parsed: %d text blocks extracted, %d chars output", len(extractor.chunks), len(result))
        return result, "html-builtin"
    except Exception as e:
        logger.error("HTML parsing failed: %s", e, exc_info=True)
        return f"[Error reading HTML: {e}]", "error"


def _parse_text(content: bytes) -> tuple[str, str]:
    """Decode plain text / markdown files."""
    try:
        text = content.decode("utf-8", errors="replace")
        logger.info("Text file decoded: %d chars", len(text))
        return text, "text-builtin"
    except Exception as e:
        logger.error("Text decoding failed: %s", e, exc_info=True)
        return f"[Error reading file: {e}]", "error"


# ─── Parser Routing ───────────────────────────────────────────────────────────

_PARSER_MAP = {
    ".docx": _parse_docx,
    ".csv": _parse_csv,
    ".json": _parse_json,
    ".html": _parse_html,
    ".txt": _parse_text,
    ".md": _parse_text,
}


# ─── Public API ───────────────────────────────────────────────────────────────

class FileParser:
    """
    Public entry point for document parsing.
    Validates input, routes to the correct parser, and returns structured output.
    """

    async def parse(self, filename: str, content: bytes, file_size: int) -> dict:
        """
        Parse file content and return extracted text with metadata.

        Returns:
            {
                "text": str,
                "filename": str,
                "extension": str,
                "character_count": int,
                "parser_used": str,
            }
        """
        ext = Path(filename).suffix.lower()
        logger.info("Parsing file '%s' (ext=%s, size=%d bytes)", filename, ext, file_size)

        # Validate
        _validate_file_size(file_size)
        _validate_extension(ext)

        # Route to parser
        if ext in DOCLING_EXTENSIONS:
            text, parser_used = _parse_with_docling(content, ext)
        elif ext in _PARSER_MAP:
            text, parser_used = _PARSER_MAP[ext](content)
        else:
            # Shouldn't reach here after validation, but safety net
            logger.error("No parser mapped for validated extension '%s'", ext)
            text, parser_used = f"[No parser for {ext}]", "error"

        logger.info(
            "File '%s' parsed — parser=%s, output=%d chars",
            filename, parser_used, len(text),
        )

        return {
            "filename": filename,
            "extension": ext,
            "text": text,
            "character_count": len(text),
            "parser_used": parser_used,
        }
